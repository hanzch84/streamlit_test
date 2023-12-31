import streamlit as st
import cv2
import easyocr
import qrcode
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from io import BytesIO
from pyzbar.pyzbar import decode
from PIL import ImageFont, ImageDraw, Image


# 자동 줄바꿈을 위한 CSS 스타일 추가

st.set_page_config(layout="centered")
st.markdown(
    """
    <style>
    div.row-widget.stRadio > div{flex-direction:row;}
    code {
        overflow: auto;
        white-space: pre-wrap;}
    </style>
    """,
    unsafe_allow_html=True)

# Calculate the table width to fit A4 paper
table_width = Cm(21.0 - 2.0)  # A4 width minus 20mm left and right margins
@st.cache_data
def maketextgrid(text_area_input):
    single_marks = [".",",","?","!"]
    # 2차원 리스트를 생성합니다.
    text_array_2D = []
    # 3개의 세로 행을 만듭니다.
    for i in range(7):
        row = ["", "", "", "", "", "", "", "", "", "", "", "", ""]
        text_array_2D.append(row)
            
    for row_num in range(7):
        for col_num in range(13):
            if text_area_input!="":
                first_char = text_area_input[0]
                text_area_input = text_area_input[1:]
                if first_char == '\n':
                    first_char = ""
                    text_area_input = (" " * (12-col_num)) + text_area_input
                elif first_char in single_marks:
                    first_char += " "
                    if text_area_input!="":
                        if text_area_input[0] == " ":
                            text_area_input = text_area_input[1:]
            else: first_char = " "
            text_array_2D[row_num][col_num] = first_char
    return text_array_2D

# 새 MS워드 문서를 만듭니다.
@st.cache_data
def create_word_document(text_grid, qr_code_data=None):
    doc = Document()
    
    # 용지설정을 가로로
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # 열과 행 및 스타일 설정
    num_rows = 7
    num_cols = 13
    style = doc.styles['Normal']
    style.font.name = '궁서체'    
    style.font.size = Pt(32)
    style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '궁서체') # 한글 폰트를 따로 설정해 준다

    # 정사각형 셀을 설정
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = doc.styles['Table Grid']

    # Set the alignment of the cells to center both horizontally and vertically
    for row in table.rows:
        row.height = Cm(1.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(1.6)  # Make cells square
            cell.height = Cm(2)

    # Set the font size and name for all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(32)  # 24-point font size
                    run.font.name = '궁서체'  # Palace font

    # Iterate over each cell in the table and insert characters from the input text grid
    for row_num in range(num_rows):
        for col_num in range(num_cols):
            cell = table.cell(row_num, col_num)
            cell.text = text_grid[row_num][col_num]
    
    doc.add_page_break()



    # Create a table with square cells
    table2 = doc.add_table(rows=num_rows, cols=num_cols)
    table2.style = doc.styles['Table Grid']

    # Set the alignment of the cells to center both horizontally and vertically
    for row in table2.rows:
        row.height = Cm(1.6)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(1.6)  # Make cells square
            cell.height = Cm(2)

    # Set the font size and name for all cells
    for row in table2.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(32)  # 24-point font size
                    run.font.name = '궁서체'  # Palace font


    # Insert QR code if QR code data is provided
    if qr_code_data:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(qr_code_data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img_bytes = BytesIO()
        img.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        doc.add_picture(img_bytes, width=Cm(4), height=Cm(4))

    # Save the Word document to a BytesIO object
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes



def ocr_label(image_data, lang, font_color):
    font_path = 'pages/D2Coding-Ver1.3.2-20180524.ttf'
    font = ImageFont.truetype(font_path, 24, encoding='utf-8')
    reader = easyocr.Reader(lang)

    # 이미지 데이터가 NumPy 배열인 경우, 해당 배열을 사용
    if isinstance(image_data, np.ndarray):
        image = image_data
    else:
        try:
            # 파일 포인터를 시작 부분으로 이동
            image_data.seek(0)
            # 파일 객체에서 데이터 읽기
            buffer = image_data.read()
        except AttributeError:
            # 이미 바이트 배열인 경우, 직접 사용
            buffer = image_data

        # 바이트 배열을 NumPy 배열로 변환
        nparr = np.frombuffer(buffer, np.uint8)
        if nparr.size == 0:
            # 배열이 비어 있으면 오류 메시지 출력
            raise ValueError("이미지 데이터가 비어 있습니다.")

        # NumPy 배열을 이미지로 디코드
        image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if image is None:
            # 이미지 디코딩 실패 시 오류 메시지 출력
            raise ValueError("이미지 디코딩에 실패했습니다.")

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    result = reader.readtext(gray)

    for (bbox, text, prob) in result:
        (top_left, top_right, bottom_right, bottom_left) = bbox
        top_left = tuple(map(int, top_left))
        bottom_right = tuple(map(int, bottom_right))
        cv2.rectangle(image, top_left, bottom_right, (0, 255, 0), 2)

        image_pil = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(image_pil)
        draw.text((top_left[0], top_left[1] - 24), text, font=font, fill=font_color)
        image = cv2.cvtColor(np.array(image_pil), cv2.COLOR_RGB2BGR)

    st.image(image, channels="BGR")
    return result

def hex_to_rgb(hex_color):
    # 앞의 '#' 기호를 제거합니다.
    hex_color = hex_color.lstrip('#')
    
    # 문자열을 세 부분으로 나누고 각각을 16진수에서 10진수로 변환합니다.
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    return (r, g, b)

def find_and_split_grid(image, grid_size=(13, 7)):
    # 이미지를 그레이스케일로 변환
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)

    # 엣지 검출
    edges = cv2.Canny(blurred, 50, 150, apertureSize=3)

    # 선 검출
    lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=100, maxLineGap=10)

    if lines is None:
        return []

    # 선을 수평, 수직으로 분류
    horizontal_lines = []
    vertical_lines = []
    for line in lines:
        x1, y1, x2, y2 = line[0]
        if abs(y2 - y1) < abs(x2 - x1): # 수평선
            horizontal_lines.append((x1, y1, x2, y2))
        else: # 수직선
            vertical_lines.append((x1, y1, x2, y2))

    # 교차점 찾기
    intersections = []
    for h_line in horizontal_lines:
        for v_line in vertical_lines:
            px, py = line_intersection(h_line, v_line)
            intersections.append((px, py))

    # 교차점을 기준으로 셀 위치 계산
    intersections = np.array(intersections)
    intersections = intersections[np.lexsort((intersections[:, 0], intersections[:, 1]))]
    #st.dataframe(intersections)
    grid_cells = get_grid_cells(intersections, grid_size)
    #st.dataframe(grid_cells)
    


    # 셀 이미지 분할
    cell_images = []
    for cell in grid_cells:
        x1, y1, x2, y2 = cell
        cell_img = image[y1:y2, x1:x2]
        cell_images.append(cell_img)

    return cell_images,horizontal_lines,vertical_lines

def line_intersection(line1, line2):
    x1, y1, x2, y2 = line1
    x3, y3, x4, y4 = line2

    d = (x1 - x2) * (y3 - y4) - (y1 - y2) * (x3 - x4)
    if d:
        px = ((x1 * y2 - y1 * x2) * (x3 - x4) - (x1 - x2) * (x3 * y4 - y3 * x4)) / d
        py = ((x1 * y2 - y1 * x2) * (y3 - y4) - (y1 - y2) * (x3 * y4 - y3 * x4)) / d
        return int(px), int(py)
    else:
        return None

def get_grid_cells(intersections, grid_size):
    expected_intersections = (grid_size[0] + 1) * (grid_size[1] + 1)
    if len(intersections) < expected_intersections:
        return []    
        raise ValueError(f"충분한 교차점이 검출되지 않았습니다. 검출된 교차점의 수: {len(intersections)}, 필요한 교차점의 수: {expected_intersections}")
        
    

    # 교차점을 x와 y 좌표별로 정렬
    intersections = sorted(intersections, key=lambda x: (x[1], x[0]))

    # 그리드 셀의 위치를 저장할 리스트
    grid_cells = []

    # 그리드의 각 셀 위치 계산
    for i in range(grid_size[1]):
        for j in range(grid_size[0]):
            top_left = intersections[i * (grid_size[0] + 1) + j]
            bottom_right = intersections[(i + 1) * (grid_size[0] + 1) + (j + 1)]
            # 셀 위치: (x1, y1, x2, y2)
            cell_position = (top_left[0], top_left[1], bottom_right[0], bottom_right[1])
            grid_cells.append(cell_position)

    return grid_cells


    
def find_grid(picture):
    try:
        # 파일 포인터를 시작 부분으로 이동
        picture.seek(0)
        # 파일 객체에서 데이터 읽기
        file_bytes = np.asarray(bytearray(picture.read()), dtype=np.uint8)
    except Exception as e:
        raise ValueError(f"파일 읽기 오류: {e}")

    # 파일 데이터로부터 이미지 디코딩
    image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    if image is None:
        raise ValueError("이미지 디코딩 실패")

    # 이미지 객체를 사용하여 그리드 셀 이미지를 얻음
    grid_cells,vertical_lines,horizontal_lines = find_and_split_grid(image)

    return grid_cells,vertical_lines,horizontal_lines

def recognize_characters(grid_cells):
    reader = easyocr.Reader(['en','ko'])  # 여기서 언어 설정을 할 수 있음
    recognized_chars = []
    for cell in grid_cells:
        if cell.size == 0 or cell is None:  # 비어 있는 셀 확인
            recognized_chars.append("")  # 비어 있는 셀에 대해서는 빈 문자열 추가
            continue

        # EasyOCR을 사용하여 셀 내 문자 인식
        result = reader.readtext(cell)
        recognized_text = result[0][-2] if result else ''
        recognized_chars.append(recognized_text)

    return recognized_chars

def draw_grid_on_image(image, horizontal_lines, vertical_lines):
    # 각 수평선을 이미지에 그립니다.
    for line in horizontal_lines:
        x1, y1, x2, y2 = line
        cv2.line(image, (x1, y1), (x2, y2), (0, 255, 0), 2)

    # 각 수직선을 이미지에 그립니다.
    for line in vertical_lines:
        x1, y1, x2, y2 = line
        cv2.line(image, (x1, y1), (x2, y2), (0, 255, 0), 2)

    return image

def on_chaange_text():
    st.session_state['first_load'] = False  # 표시되지 않도록 설정합니다.


# 세션 상태에 'first_load' 키가 없으면 True를 설정합니다. (처음 로딩 시)
if 'first_load' not in st.session_state:
    st.session_state['first_load'] = True

# Streamlit app
st.title("도전! 예쁜 글씨 쓰기👍")
st.subheader("글씨쓰기 학습지 만들기 + 자동 채점하기")
st.write("학습지의 앞면은 따라쓰기, 뒷면은 빈 칸입니다. 뒷면에 있는 qr코드로 손글씨의 정답을 인식합니다.")

# 첫 로딩 시에만 텍스트 영역에 기본값을 표시합니다.
if st.session_state['first_load']:
    text_to_insert = st.text_area("Enter text (13x7 characters):", height=100,on_change=on_chaange_text, value='오늘도 또 우리 수탉이 막 쫓기었다. 내가 점심을 먹고 나무를 하러 갈 양으로 나올 때이었다. 산으로 올라서려니까 등뒤에서 푸드득 푸드득 하고 닭의 횃소리가 야단이다.')
else:
    text_to_insert = st.text_area("Enter text (13x7 characters):", height=100)

result_matrix = maketextgrid(text_to_insert)
df =pd.DataFrame(result_matrix)          
st.dataframe(df,hide_index=True,use_container_width=True)
col_make_btn, col_down_btn = st.columns(2)
if col_make_btn.button("MS WORD 문서 생성하기"):
    
    # Create the Word document
    result_qr = ''.join([''.join(row) for row in result_matrix])
    doc_bytes = create_word_document(result_matrix, result_qr)

    # Offer the document for download
    col_down_btn.download_button(
        label="MS WORD 문서 다운로드",
        data=doc_bytes,
        key="word_doc",
        file_name="table_with_text.docx",
        mime="application/octet-stream",
    )
languages_selected = ["ko", "en"]
col_color_pick,col_color_label, col_image_pick = st.columns([1,4,4])

font_color = hex_to_rgb(col_color_pick.color_picker('폰트 색상을 지정하세요.','#00FF00',label_visibility='collapsed'))
col_color_label.write('폰트 색상을 지정하세요.')
radio_cam_option = col_image_pick.radio("카메라 촬영 vs 파일 업로드", ["카메라 촬영", "파일 업로드"],label_visibility='collapsed')

if radio_cam_option == "카메라 촬영":
    picture = st.camera_input("#사진을 찍으면 문자를 인식해요!")
else:
    picture = st.file_uploader('이미지를 업로드 하세요.', type=['png', 'jpg', 'jpeg'])


try:
        
    if picture is not None:
        outputs = ocr_label(picture, languages_selected,font_color)
        if outputs.__len__() != 0:
            ocr_text = list(zip(*outputs))[1]
            st.write(ocr_text)
            d = decode(Image.open(picture))
        for data in d: #QR출력
            #st.code(f"{data.type} = {data.data.decode('utf-8')}",language="ada")
            #st.warning(f"{data.type} = {data.data.decode('utf-8')}")
            st.write(f"{data.type} = {data.data.decode('utf-8')}")

        # 그리드 찾기 및 캐릭터 인식
        grid_cells, horizontal_lines, vertical_lines = find_grid(picture)
        recognized_chars = recognize_characters(grid_cells)
        np_array = np.array(recognized_chars)
        #st.table(np_array.reshape(7, 13))

        # 파일 포인터를 시작 부분으로 이동
        picture.seek(0)
        # 파일 객체에서 데이터 읽기
        file_bytes = np.asarray(bytearray(picture.read()), dtype=np.uint8)

        # 파일 데이터로부터 이미지 디코딩
        image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        
        # 그리드 찾기 및 이미지에 그리드 그리기
        image_with_grid = draw_grid_on_image(image.copy(), horizontal_lines, vertical_lines)

        # Streamlit을 통해 이미지 출력
        st. subheader("그리드 인식")
        st.image(image_with_grid, channels="BGR")

except:
    pass